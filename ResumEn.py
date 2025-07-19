import argparse
from docx import Document
from docx.shared import Pt
from openai import OpenAI

def main():
    #provide user with help about the tool and how to use it
    parser = argparse.ArgumentParser(
        prog="ResumEn.py",
        description="Modifies your resume based on job descriptions to combat ATS filtering.",
        usage="ResumEn.py [arguments] [file]"
    )

    #accepted command-line arguments
    parser.add_argument("-f","--file", help="specifies the path to your resume (.docx only supported)", required=True)
    parser.add_argument("-i", "--input", help="path of input file containing job descriptions", required=True)
    parser.add_argument("-q", "--qualifier", help="path of the qualifier file for generation of professional summary.", required=True)

    #enables argparse; displays/prints to command line
    args = parser.parse_args()

    #get skills
    skills = extractSkills(args.input)

    #write skills to file
    skills2Docx(skills)

    #generate summary, returns summary
    summary = summarize(args.qualifier, args.input)

    #write summary to file
    sum2Docx(args.file, summary)

#generates a professional summary and returns it
def summarize(qualPath, descPath):
    #opens path to qualifiers file
    qualifiers = fetchContentsFromFile(qualPath)
        
    #opens path to job description file
    description = fetchContentsFromFile(descPath)
    
    # OPEN AI API KEY GOES HERE
    gpt = OpenAI(api_key="")

    log("Generating a professional summary...")

    response = gpt.responses.create(
        model="gpt-4-turbo-2024-04-09",
        input=f"In 500 characters or less, (do not exceed this value) create a professional summary for use on a resume tailored for the attached job description. You should be using more keywords from the below job description to optimize for ATS, and less of my qualifications. Use my qualifications SPARRINGLY. Again, You do not need to force every qualifications listed in the summary. This should read natural like a professional summary would, but as one optimized for ATSs. Do not to use the words 'adept' or 'poised' in this summary. The summary is attached below, followed by the job description. Give precedence to the most relevant skills. ONLY INCLUDE THE SUMMARY. I just want the summary. Do not preface the summary by saying it is a summary. Thank You!\nMy skills:\n{qualifiers}\nJob Description:\n{description}"
    )

    return response.output_text

def sum2Docx(path, summary):
    doc = Document(path)

    style = doc.styles['Normal']
    
    #specifies font to match resume
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    #location of job summary paragraph
    starting_index = 4

    #deletes previous job summary paragraph
    for _ in range(1):
        p = doc.paragraphs[starting_index]._element
        p.getparent().remove(p)
        break
    
    #place new line in place of deleted paragraph
    doc.paragraphs[starting_index]._insert_paragraph_before()

    #write paragraph to new line
    doc.paragraphs[starting_index].text = summary

    doc.save("tResume.docx")

    log("Summary successfully generated and written to file.")    

#takes job description from input file; extracts skills.
def extractSkills(path):
    #opens path to job description file    
    description = fetchContentsFromFile(path)

    # OPEN AI API KEY GOES HERE
    gpt = OpenAI(api_key="")

    log("Generating skills...")

    #specifies openai chatgpt prompt and expected output
    response = gpt.responses.create(
        model="gpt-4.1-2025-04-14",
        input=f"Extract exactly four of the most important skills from this job description and insert them in chronological order. Do not include any other text, and separate each skill with a comma.:\n{description}",
    )

    #gets skills, adds each element separated by commas into array
    skills = str.split(response.output_text, ', ')

    #return skills array to write later to .docx resume
    return skills 

#takes generated skills and writes them to .docx
def skills2Docx(skills):
    path = ("tResume.docx")

    doc = Document(path)
    style = doc.styles['Normal']
    
    #specifies font to match resume
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    #finds 'Skills' header
    for index, paragraph in enumerate(doc.paragraphs):
        if 'PROFESSIONAL SKILLS' in paragraph.text:
            skillsLine = index    

    #applies style changes to the soon-to-be 
    #added text instead of the default (Arial).
    paragraph.style = doc.styles['Normal']

    #writes skills list to resume
    for i in range(4):
        skillsLine += 1
        doc.paragraphs[skillsLine].text = skills[i]
        doc.save('tResume.docx')

    log("Skills successfully generated and written to file.")

def log(message: str):
    print(f"\n[!]\t{message}")

def fetchContentsFromFile(path):
        #opens path to job description file
    with open(path, "r") as file:
        contents = file.read()

    return contents

if __name__=="__main__":
    main()